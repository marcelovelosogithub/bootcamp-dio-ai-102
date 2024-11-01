# import json
import os
import uuid
from io import BytesIO

import requests
import streamlit as st
from docx import Document
from dotenv import load_dotenv

# Carregar variáveis de ambiente do arquivo .env
load_dotenv()

API_KEY = os.getenv("TRANSLATOR_API_KEY")
ENDPOINT = os.getenv("TRANSLATOR_ENDPOINT")
LOCATION = os.getenv("TRANSLATOR_LOCATION")


# Função para realizar a tradução
def traduzir_texto(texto, idioma_origem, idioma_destino):
    path = "/translate"
    constructed_url = ENDPOINT + path

    params = {"api-version": "3.0", "from": idioma_origem, "to": [idioma_destino]}

    headers = {
        "Ocp-Apim-Subscription-Key": API_KEY,
        "Ocp-Apim-Subscription-Region": LOCATION,
        "Content-type": "application/json",
        "X-ClientTraceId": str(uuid.uuid4()),
    }

    body = [{"text": texto}]

    # Fazendo a requisição de tradução
    response = requests.post(constructed_url, params=params, headers=headers, json=body)

    # Tratamento de resposta
    if response.status_code == 200:
        traduzido = response.json()[0]["translations"][0]["text"]
        return traduzido
    else:
        return f"Erro na tradução: {response.status_code}"


# Função para carregar o texto do arquivo Word
def carregar_texto_word(arquivo):
    doc = Document(arquivo)
    texto_completo = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
    return texto_completo


# Função para criar o arquivo traduzido para download
def criar_arquivo_para_download(texto, nome_arquivo="tradução.docx"):
    doc = Document()
    for linha in texto.split("\n"):
        doc.add_paragraph(linha)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Interface do usuário com Streamlit
st.title("Tradutor de Arquivos Word")

# Carregar arquivo Word
arquivo = st.file_uploader("Carregar arquivo Word", type=["docx"])

# Seleção de idiomas
idiomas = ["en", "fr", "es", "de", "pt", "it"]  # Pode adicionar mais idiomas
idioma_origem = st.selectbox("Idioma de origem", idiomas)
idioma_destino = st.selectbox("Idioma de destino", idiomas)

# Botão para traduzir
if st.button("Traduzir"):
    if arquivo is not None:
        # Carregar e exibir o texto original
        texto_original = carregar_texto_word(arquivo)
        st.markdown("### Texto Original:")
        st.text_area("Texto Original", texto_original, height=200)

        # Realizar tradução
        traducao = traduzir_texto(texto_original, idioma_origem, idioma_destino)

        # Exibir o resultado da tradução
        st.markdown("### Tradução:")
        st.text_area("Tradução", traducao, height=200)

        # Botão de download da tradução
        buffer = criar_arquivo_para_download(traducao)
        st.download_button(
            label="Baixar Tradução",
            data=buffer,
            file_name="tradução.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.error("Por favor, carregue um arquivo Word.")
